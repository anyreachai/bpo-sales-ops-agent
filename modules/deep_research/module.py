from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from modules._base import BaseModule
from modules.deep_research.prompts import build_research_prompt
from orchestrator.config import settings
from shared.anthropic_client import call_opus_with_search
from shared.storage import artifact_path
from shared.types import Artifact, ModuleResult, SessionContext

logger = logging.getLogger(__name__)

# -- Color palette --
NAVY = RGBColor(0x1A, 0x1F, 0x3D)
BLUE_ACCENT = RGBColor(0x2B, 0x5C, 0x8A)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# =====================================================================
# markdown_to_docx helper
# =====================================================================


def markdown_to_docx(
    markdown_text: str,
    output_path: Path,
    company_name: str,
) -> Path:
    """Convert a markdown research report into a styled .docx file.

    Handles:
    - Heading levels (# ## ###)
    - Bullet lists (- or * prefix)
    - Markdown tables (| delimited)
    - Bold and italic inline formatting
    - A cover page and header/footer with page numbers
    """
    doc = Document()

    # ---- Global font defaults ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Configure heading styles
    _configure_heading_styles(doc)

    # ---- Cover page ----
    _add_cover_page(doc, company_name)

    # ---- Parse markdown into document body ----
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip blank lines
        if not stripped:
            i += 1
            continue

        # Heading detection
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Strip any trailing # characters
            heading_text = re.sub(r"\s*#+\s*$", "", heading_text)
            # Strip bold markers from heading text
            heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text)
            p = doc.add_heading(heading_text, level=min(level, 4))
            # Ensure heading font is Calibri and navy
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.color.rgb = NAVY
            i += 1
            continue

        # Table detection (line starts with |)
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue

        # Bullet list detection
        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list detection
        num_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if num_match:
            text = num_match.group(1)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraph — accumulate consecutive non-special lines
        para_lines = []
        while i < len(lines):
            current = lines[i].strip()
            if not current:
                i += 1
                break
            if (
                re.match(r"^#{1,4}\s+", current)
                or current.startswith("|")
                or re.match(r"^[-*]\s+", current)
                or re.match(r"^\d+[.)]\s+", current)
            ):
                break
            para_lines.append(current)
            i += 1

        if para_lines:
            p = doc.add_paragraph()
            _add_formatted_text(p, " ".join(para_lines))

    # ---- Header and footer ----
    _add_header_footer(doc, company_name)

    # ---- Save ----
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# =====================================================================
# Internal helpers
# =====================================================================


def _configure_heading_styles(doc: Document) -> None:
    """Set up heading styles with Calibri + navy color."""
    for level, size in [(1, 20), (2, 16), (3, 13), (4, 11)]:
        style_name = f"Heading {level}"
        try:
            h_style = doc.styles[style_name]
        except KeyError:
            continue
        h_style.font.name = "Calibri"
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = NAVY
        h_style.font.bold = True
        # Space before/after
        h_style.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
        h_style.paragraph_format.space_after = Pt(4)


def _add_cover_page(doc: Document, company_name: str) -> None:
    """Insert a simple cover page with company name, report title, and date."""
    # Add some blank space at the top
    for _ in range(6):
        doc.add_paragraph()

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name)
    run.font.name = "Calibri"
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.bold = True

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = BLUE_ACCENT
    run.font.size = Pt(10)

    # Report title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Deep Research Report")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE_ACCENT

    # Spacer
    doc.add_paragraph()

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GRAY

    # Spacer
    doc.add_paragraph()

    # Prepared by
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by Anyreach, Inc.")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY
    run.italic = True

    # Page break after cover
    doc.add_page_break()


def _add_table(doc: Document, table_lines: list[str]) -> None:
    """Parse markdown table lines and add a styled Word table.

    Expected format:
        | Header1 | Header2 | Header3 |
        |---------|---------|---------|
        | Cell1   | Cell2   | Cell3   |
    """
    if len(table_lines) < 2:
        # Not a valid table, just dump as paragraphs
        for line in table_lines:
            doc.add_paragraph(line)
        return

    def _parse_row(line: str) -> list[str]:
        """Split a markdown table row into cell strings."""
        # Strip leading/trailing |
        inner = line.strip().strip("|")
        return [cell.strip() for cell in inner.split("|")]

    # Detect and skip separator row (e.g., |---|---|)
    rows: list[list[str]] = []
    for line in table_lines:
        if re.match(r"^\|[\s\-:]+\|$", line.strip().replace("|", "|").strip()):
            # Check if every cell is only dashes/colons/spaces
            cells = _parse_row(line)
            if all(re.match(r"^[\s\-:]*$", c) for c in cells):
                continue
        rows.append(_parse_row(line))

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    # Normalize column count
    for r in rows:
        while len(r) < num_cols:
            r.append("")

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = "Calibri"
            run.font.size = Pt(9)

            if row_idx == 0:
                # Header row styling
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
                # Set cell background to navy
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(
                    qn("w:shd"),
                    {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): "1A1F3D",
                    },
                )
                shading.append(shading_elem)
            else:
                run.font.color.rgb = DARK_GRAY
                # Alternate row shading
                if row_idx % 2 == 0:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elem = shading.makeelement(
                        qn("w:shd"),
                        {
                            qn("w:val"): "clear",
                            qn("w:color"): "auto",
                            qn("w:fill"): "F2F2F2",
                        },
                    )
                    shading.append(shading_elem)

    # Small space after table
    doc.add_paragraph()


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text to a paragraph, handling **bold** and *italic* markers."""
    # Pattern: split on **bold** and *italic* markers
    # Process bold first, then italic within each segment
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Bold segment — strip markers and check for nested italic
            inner = part[2:-2]
            _add_italic_segments(paragraph, inner, bold=True)
        else:
            _add_italic_segments(paragraph, part, bold=False)


def _add_italic_segments(paragraph, text: str, bold: bool = False) -> None:
    """Handle *italic* within a text segment."""
    parts = re.split(r"(\*[^*]+?\*)", text)
    for part in parts:
        if part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY
            run.italic = True
            run.bold = bold
        elif part:
            run = paragraph.add_run(part)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY
            run.bold = bold


def _add_header_footer(doc: Document, company_name: str) -> None:
    """Add a header with the company name and footer with page numbers."""
    for section in doc.sections:
        # Different first page (cover page gets no header/footer)
        section.different_first_page_header_footer = True

        # Header
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        hp.text = ""
        run = hp.add_run(f"{company_name}  |  Deep Research Report")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = BLUE_ACCENT
        run.italic = True

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = ""

        # "Prepared by Anyreach, Inc." on the left side
        run = fp.add_run("Prepared by Anyreach, Inc.   ")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = DARK_GRAY
        run.italic = True

        # Page number field
        run2 = fp.add_run()
        fld_char_begin = run2._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run2._element.append(fld_char_begin)

        run3 = fp.add_run()
        instr_text = run3._element.makeelement(qn("w:instrText"), {})
        instr_text.text = " PAGE "
        run3._element.append(instr_text)
        run3.font.name = "Calibri"
        run3.font.size = Pt(8)
        run3.font.color.rgb = DARK_GRAY

        run4 = fp.add_run()
        fld_char_end = run4._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run4._element.append(fld_char_end)


# =====================================================================
# DeepResearchModule
# =====================================================================


class DeepResearchModule(BaseModule):
    name = "deep_research"

    def should_run(self, ctx: SessionContext) -> bool:
        return "deep_research" in ctx.deliverables_requested

    async def run(self, ctx: SessionContext) -> ModuleResult:
        company = ctx.target_company or "Unknown Company"
        logger.info("Starting deep research for %s", company)

        # 1. Build the prompt
        intake_dict = None
        if ctx.intake:
            intake_dict = {
                "contact_name": ctx.intake.contact_name,
                "contact_title": ctx.intake.contact_title,
                "target_business_area": ctx.intake.target_business_area,
                "pain_points": ctx.intake.pain_points,
                "current_setup": ctx.intake.current_setup,
            }

        bpo_context = ""
        if ctx.bpo:
            bpo_context = (
                f"Requesting BPO partner: {ctx.bpo.name} "
                f"(key contacts: {', '.join(ctx.bpo.key_contacts) if ctx.bpo.key_contacts else 'N/A'})"
            )

        prompt = build_research_prompt(
            company_name=company,
            company_url=ctx.target_url,
            intake=intake_dict,
            bpo_context=bpo_context,
        )

        system = (
            "You are a corporate intelligence analyst with deep expertise in BPO, "
            "contact center operations, and customer experience outsourcing. "
            "Produce thorough, evidence-backed research reports. Use web search "
            "extensively to gather current information. Be candid and analytical."
        )

        # 2. Call Claude Opus with web search
        logger.info("Calling Claude Opus with web search for %s", company)
        markdown_response = await call_opus_with_search(
            api_key=settings.ANTHROPIC_API_KEY,
            prompt=prompt,
            system=system,
            max_tokens=16000,
            thinking_budget=10000,
        )

        if not markdown_response or len(markdown_response.strip()) < 200:
            logger.error("Research response too short or empty for %s", company)
            return ModuleResult(
                module_name=self.name,
                status="failed",
                error="Research response was too short or empty",
                metadata={"response_length": len(markdown_response) if markdown_response else 0},
            )

        logger.info(
            "Received research response for %s (%d chars)",
            company,
            len(markdown_response),
        )

        # 3. Generate .docx
        docx_path = artifact_path(
            session_id=ctx.session_id,
            company=company,
            suffix="Deep_Research",
            ext="docx",
        )

        try:
            markdown_to_docx(markdown_response, docx_path, company)
        except Exception as e:
            logger.exception("Failed to generate .docx for %s", company)
            return ModuleResult(
                module_name=self.name,
                status="failed",
                error=f"DOCX generation failed: {e}",
                metadata={"response_length": len(markdown_response)},
            )

        file_size = docx_path.stat().st_size
        logger.info("Generated %s (%d bytes)", docx_path.name, file_size)

        # 4. Build artifact
        artifact = Artifact(
            filename=docx_path.name,
            path=docx_path,
            artifact_type="deep_research",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size_bytes=file_size,
        )

        return ModuleResult(
            module_name=self.name,
            status="success",
            artifacts=[artifact],
            metadata={
                "company": company,
                "url": ctx.target_url,
                "response_length": len(markdown_response),
                "docx_size_bytes": file_size,
                "docx_path": str(docx_path),
            },
        )
